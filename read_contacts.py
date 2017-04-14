from docx import Document
import pprint


# def read_docx(input_file):
#     doc = Document(input_file)
#
#     table = doc.tables[0]
#
#     cont = list()
#
#     department = str()
#
#     for row in table.rows:
#
#         if row.cells[0].text == row.cells[1].text:
#             department = row.cells[0].text.strip(' ')
#             print(department)
#
#         else:
#             unit = dict()
#
#             for i, cell in enumerate(row.cells, 1):
#
#                 if 'IT (технический)' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'IT'
#
#                 elif 'VIP-клиентами' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'VIP'
#
#                 elif 'Департамент элитной недвижимости' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'ДЭН'
#
#                 elif 'Департамент долевого строительства' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'Долевое'
#
#                 elif 'Департамент аренды элитной' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'Аренда'
#
#                 elif 'Степанова Михаила' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'Отдел Степанова'
#
#                 elif 'CALL – центр' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'CALL – центр'
#
#                 elif 'Brilliant Home' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'Brilliant Home'
#
#                 elif 'Независимые агенты' in department:
#                     unit['Organization 1 - Name'] = 'Мир Квартир'
#                     unit['Organization 1 - Title'] = 'Независимые агенты'
#
#                 elif 'Мир Квартир-Элит' in department:
#                     unit['Organization 1 - Name'] = 'МК Элит'
#
#                 elif 'отдел Вторичной недвижимости' in department:
#                     unit['Organization 1 - Name'] = 'МК Элит'
#                     unit['Organization 1 - Title'] = 'Вторичка'
#
#                 elif 'Зубенко Артема' in department:
#                     unit['Organization 1 - Name'] = 'МК Элит'
#                     unit['Organization 1 - Title'] = 'Зубенко Артема'
#
#                 elif 'МКЭлит' in department:
#                     unit['Organization 1 - Name'] = 'МК Элит'
#                     unit['Organization 1 - Title'] = 'ОДС'
#
#                 if i == 1:
#                     full_name = cell.text.split(' ')
#                     unit['Family Name'] = full_name[0]
#                     unit['Given Name'] = full_name[1]
#                 if i == 2:
#                     unit['Birthday'] = cell.text
#                 if i == 3:
#                     unit['Phone 1 - Value'] = cell.text.strip()
#                 if i == 4:
#                     unit['E-mail 1 - Value'] = cell.text.strip()
#
#             unit['E-mail 1 - Type'] = 'Work'
#             unit['Phone 1 - Type'] = 'Mobile'
#             unit['Name'] = ' '.join([unit['Given Name'], unit['Family Name']])
#             cont.append(unit)
#     return cont


def read_docx(input_file):
    doc = Document(input_file)

    table = doc.tables[0]

    cont = list()

    department = str()

    with open('test.vcf', 'w') as f:
        f.write('')

    for row in table.rows:

        if row.cells[0].text == row.cells[1].text:
            department = row.cells[0].text.strip(' ')
            print(department)

        else:

            with open('test.vcf', 'at') as f:

                f.write("BEGIN:VCARD \n")
                f.write("VERSION:3.0 \n")

            for i, cell in enumerate(row.cells, 1):

                if i == 1:
                    print(cell.text)
                    # full_name = cell.text.split(' ')
                    #                     f.write(''.join(['FN:', cell.text]))
                    #                     f.write(''.join('N:', ';'.join(cell.text.split(' '))))
                    #                 if i == 2:
                    #                     unit['Birthday'] = cell.text
                    #                 if i == 3:
                    #                     unit['Phone 1 - Value'] = cell.text.strip()
                    #                 if i == 4:
                    #                     unit['E-mail 1 - Value'] = cell.text.strip()
                    #
                    #
                    #                 if 'IT (технический)' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'IT'
                    #
                    #                 elif 'VIP-клиентами' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'VIP'
                    #
                    #                 elif 'Департамент элитной недвижимости' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'ДЭН'
                    #
                    #                 elif 'Департамент долевого строительства' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'Долевое'
                    #
                    #                 elif 'Департамент аренды элитной' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'Аренда'
                    #
                    #                 elif 'Степанова Михаила' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'Отдел Степанова'
                    #
                    #                 elif 'CALL – центр' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'CALL – центр'
                    #
                    #                 elif 'Brilliant Home' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'Brilliant Home'
                    #
                    #                 elif 'Независимые агенты' in department:
                    #                     unit['Organization 1 - Name'] = 'Мир Квартир'
                    #                     unit['Organization 1 - Title'] = 'Независимые агенты'
                    #
                    #                 elif 'Мир Квартир-Элит' in department:
                    #                     unit['Organization 1 - Name'] = 'МК Элит'
                    #
                    #                 elif 'отдел Вторичной недвижимости' in department:
                    #                     unit['Organization 1 - Name'] = 'МК Элит'
                    #                     unit['Organization 1 - Title'] = 'Вторичка'
                    #
                    #                 elif 'Зубенко Артема' in department:
                    #                     unit['Organization 1 - Name'] = 'МК Элит'
                    #                     unit['Organization 1 - Title'] = 'Зубенко Артема'
                    #
                    #                 elif 'МКЭлит' in department:
                    #                     unit['Organization 1 - Name'] = 'МК Элит'
                    #                     unit['Organization 1 - Title'] = 'ОДС'
                    #
                    #             unit['E-mail 1 - Type'] = 'Work'
                    #             unit['Phone 1 - Type'] = 'Mobile'
                    #             unit['Name'] = ' '.join([unit['Given Name'], unit['Family Name']])
                    #             cont.append(unit)
                    # return cont


# def generate_csv(cont):
#     with open('test.csv', 'w', encoding='utf16') as f:
#         f.write(','.join(cont[0].keys()))
#     # print(','.join(cont[0].keys()))
#     for unit in cont:
#         with open('test.csv', 'a', encoding='utf16') as f:
#             f.write(''.join([','.join(unit.values()), '\n']))
#             # print(','.join(unit.values()))

def generate_csv(cont):
    with open('test.csv', 'w', encoding='utf16') as f:
        f.write(','.join(cont[0].keys()))
    # print(','.join(cont[0].keys()))
    for unit in cont:
        with open('test.csv', 'a', encoding='utf16') as f:
            f.write(''.join([','.join(unit.values()), '\n']))
            # print(','.join(unit.values()))


def main():
    input_file = 'cont.docx'
    cont = read_docx(input_file)
    generate_csv(cont)
    print(len(cont))


if __name__ == '__main__':
    main()
